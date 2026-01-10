from docx import Document
from backend.tailor import extract_text_from_docx
import os

def test_table_extraction():
    # Setup: Create a docx with a table
    test_filename = "test_table_resume.docx"
    doc = Document()
    doc.add_paragraph("Header Text")
    
    table = doc.add_table(rows=1, cols=2)
    cell1 = table.cell(0, 0)
    cell1.text = "Table Cell 1 Content"
    
    cell2 = table.cell(0, 1)
    cell2.text = "Table Cell 2 Content"
    
    doc.save(test_filename)
    
    # Act: Extract text
    extracted_text = extract_text_from_docx(test_filename)
    print("Extracted Text:\n", extracted_text)
    
    # Assert
    assert "Header Text" in extracted_text
    assert "Table Cell 1 Content" in extracted_text
    assert "Table Cell 2 Content" in extracted_text
    
    print("SUCCESS: Table content extracted.")
    
    # Cleanup
    try:
        os.remove(test_filename)
    except:
        pass

if __name__ == "__main__":
    test_table_extraction()
