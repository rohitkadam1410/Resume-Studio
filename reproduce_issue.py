from docx import Document
from docx.shared import Pt, RGBColor
import os
from backend.tailor import apply_edits_to_docx

def create_test_docx(filename):
    doc = Document()
    p = doc.add_paragraph()
    run1 = p.add_run("This is ")
    run2 = p.add_run("bold")
    run2.bold = True
    run3 = p.add_run(" and ")
    run4 = p.add_run("italic")
    run4.italic = True
    run5 = p.add_run(" text.")
    doc.save(filename)

def check_formatting(filename):
    doc = Document(filename)
    para = doc.paragraphs[0]
    
    # We expect "This is REPLACED and italic text."
    # "REPLACED" should ideally be bold if we want to preserve the style of the replaced text.
    # Or at least, "italic" should remain italic.
    
    print(f"Full text: {para.text}")
    print("Runs:")
    for run in para.runs:
        print(f"  Text: '{run.text}', Bold: {run.bold}, Italic: {run.italic}")

def test_formatting_preservation():
    input_file = "test_format_repro.docx"
    output_file = "test_format_repro_out.docx"
    
    create_test_docx(input_file)
    print("Original File:")
    check_formatting(input_file)
    
    # Scenario 1: Replace "bold" (which is in a single run) with "REPLACED"
    # This should hit the 'replace in run' logic if exact match works
    edits = [
        {"action": "replace", "target_text": "bold", "new_content": "REPLACED"}
    ]
    
    print("\nApplying Edit 1: 'bold' -> 'REPLACED' ...")
    apply_edits_to_docx(input_file, edits, output_file)
    check_formatting(output_file)

    # Scenario 2: Replace "This is bold" (spans runs) with "New Start"
    # This will hit the fallback logic
    edits = [
        {"action": "replace", "target_text": "This is bold", "new_content": "New Start"}
    ]
    print("\nApplying Edit 2: 'This is bold' -> 'New Start' ...")
    apply_edits_to_docx(input_file, edits, output_file)
    check_formatting(output_file)

if __name__ == "__main__":
    test_formatting_preservation()
