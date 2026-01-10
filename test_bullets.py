from docx import Document
from backend.tailor import apply_edits_to_docx
import os

def test_bullet_preservation():
    test_docx = "test_bullets.docx"
    doc = Document()
    
    # Check 1: Separate Runs (Run 1 = Bullet, Run 2 = Text)
    p = doc.add_paragraph()
    run1 = p.add_run("•")
    run1.font.name = "Symbol" # Bullet font
    run2 = p.add_run(" Original Text")
    run2.font.name = "Arial"  # Text font
    
    doc.save(test_docx)
    print("Created test_bullets.docx with multi-run bullet.")
    
    # Edit: Replace "Original Text" (matches Run 2 exactly)
    edits = [{
        "target_text": "Original Text",
        "new_content": "New Tailored Content",
        "action": "replace"
    }]
    
    output_docx = "test_bullets_tailored.docx"
    apply_edits_to_docx(test_docx, edits, output_docx)
    
    # Verify
    doc_out = Document(output_docx)
    p_out = doc_out.paragraphs[0]
    
    print(f"Total Runs: {len(p_out.runs)}")
    if len(p_out.runs) >= 2:
        r1 = p_out.runs[0]
        r2 = p_out.runs[1]
        print(f"Run 1 Text: '{r1.text}', Font: {r1.font.name}")
        print(f"Run 2 Text: '{r2.text}', Font: {r2.font.name}")
        
        if r1.text == "•" and r1.font.name == "Symbol":
             print("PASS 1: Bullet run preserved.")
        else:
             print("FAIL 1: Bullet run modified.")
             
        if "New Tailored Content" in r2.text and r2.font.name == "Arial":
             print("PASS 2: Text updated and Arial font kept.")
        else:
             print("FAIL 2: Text run incorrect.")
    else:
        # If runs merged, check fallback logic
        print("Runs merged (Fallback Logic triggered?)")
        r_all = p_out.runs[0]
        print(f"Full Text: '{r_all.text}'")
        print(f"Font: {r_all.font.name}")
        
        if r_all.font.name == "Arial":
            print("PASS 3 (Fallback): Text font (Arial) applied to merged content.")
        elif r_all.font.name == "Symbol":
            print("FAIL 3: Symbol font applied to text!")

if __name__ == "__main__":
    test_bullet_preservation()
