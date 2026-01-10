from pdf2docx import Converter
from docx2pdf import convert
import pythoncom
from docx import Document
from docx.oxml.ns import qn

def sanitize_docx_layout(docx_path: str):
    """
    Removes manual page breaks, section breaks, and restrictive paragraph properties 
    to allow content to flow naturally.
    """
    doc = Document(docx_path)
    
    def clean_para(para):
        # Reset restrictive flow properties
        if para.paragraph_format.page_break_before:
            para.paragraph_format.page_break_before = False
        if para.paragraph_format.keep_with_next:
            para.paragraph_format.keep_with_next = False
        if para.paragraph_format.keep_together:
            para.paragraph_format.keep_together = False
            
        # Remove manual breaks (<w:br>) in runs
        for run in para.runs:
            if 'lastRenderedPageBreak' in run._element.xml:
                pass
            brs = run._element.findall(qn('w:br'))
            for br in brs:
                run._element.remove(br)

    # 1. Clean Main Document Paragraphs
    for para in doc.paragraphs:
        clean_para(para)
        
        # Remove Section Breaks (w:sectPr) stored in paragraph properties
        # This effectively merges sections into a continuous flow
        if para._element.pPr is not None:
            sectPr = para._element.pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                para._element.pPr.remove(sectPr)

    # 2. Clean Tables (Recursively clean paragraphs in cells)
    for table in doc.tables:
        for row in table.rows:
            # Table rows also have height properties that might cause splits
            # We can't easily reset trHeight without low-level access, usually okay.
            for cell in row.cells:
                for para in cell.paragraphs:
                    clean_para(para)
                
                # Check for nested tables? python-docx only goes one level deep easily 
                # but usually pdf2docx doesn't nest deeply.
                
    doc.save(docx_path)

def pdf_to_docx(pdf_path: str) -> str:
    docx_path = pdf_path.replace(".pdf", ".docx")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    
    # Sanitize immediately after conversion
    sanitize_docx_layout(docx_path)
    
    return docx_path

def docx_to_pdf(docx_path: str) -> str:
    # Initialize COM library for Windows
    pythoncom.CoInitialize()
    pdf_path = docx_path.replace(".docx", ".pdf")
    
    # Ensure fresh export by deleting existing file if any
    import os
    if os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
        except:
            pass
            
    convert(docx_path, pdf_path)
    return pdf_path

