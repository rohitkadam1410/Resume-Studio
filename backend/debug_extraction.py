from docx import Document
from docx.oxml.ns import qn
import os

def debug_extraction_sources(docx_path):
    doc = Document(docx_path)
    output = []
    
    output.append(f"Analyzing File: {docx_path}\n")
    
    # 1. Headers
    output.append("--- HEADERS ---")
    for section in doc.sections:
        for p in section.header.paragraphs:
            if p.text.strip():
                output.append(f"[Header P] {p.text}")
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                             output.append(f"[Header T] {p.text}")
                             
    # 2. Paragraphs
    output.append("\n--- BODY PARAGRAPHS ---")
    for p in doc.paragraphs:
        if p.text.strip():
            output.append(f"[Body P] {p.text}")
            
    # 3. Tables
    output.append("\n--- BODY TABLES ---")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        output.append(f"[Body T] {p.text}")

    # 4. Textboxes
    output.append("\n--- TEXTBOXES ---")
    for element in doc.element.body.iter():
       if element.tag.endswith('txbxContent'):
           for p in element.findall(qn('w:p')):
               text = ""
               for r in p.findall(qn('w:r')):
                   for t in r.findall(qn('w:t')):
                       if t.text: text += t.text
               if text.strip():
                   output.append(f"[Textbox] {text}")

    return "\n".join(output)

def debug_extraction():
    # Find a temp docx file
    files = [f for f in os.listdir(".") if f.endswith(".docx") and f.startswith("temp_") and not f.endswith("_tailored.docx")]
    if not files:
        print("No temp docx files found.")
        return

    target_file = files[0]
    
    try:
        text = debug_extraction_sources(target_file)
        
        with open("debug_output_detailed.txt", "w", encoding="utf-8") as f:
            f.write(text)
            
        print(f"Extraction analysis complete. Output written to debug_output_detailed.txt.")
        
    except Exception as e:
        print(f"Extraction failed: {e}")

if __name__ == "__main__":
    debug_extraction()
