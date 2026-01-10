import os
import sys
sys.path.append('.')
from tailor import extract_text_from_docx
from docx import Document

def test_full_extraction():
    # Find a temp docx file
    files = [f for f in os.listdir(".") if f.endswith(".docx") and f.startswith("temp_") and not f.endswith("_tailored.docx")]
    if not files:
        print("No temp docx files found.")
        return

    target_file = files[0]
    print(f"Testing extraction on: {target_file}")
    
    doc = Document(target_file)
    print(f"Document has {len(doc.paragraphs)} body paragraphs and {len(doc.tables)} body tables.")
    
    text = extract_text_from_docx(target_file)
    
    with open("check_output.txt", "w", encoding="utf-8") as f:
        f.write(text)
    
    print(f"Extracted {len(text)} characters to check_output.txt")
    
    # Check for specific fragments again
    fragments = ["process efficiency", "time for regulatory professionals", "PROFESSIONAL EXPERIENCE"]
    print("\n--- Fragment Check ---")
    for frag in fragments:
        if frag.lower() in text.lower():
            print(f"[x] Found: '{frag}'")
        else:
            # Try without spaces too
            if frag.replace(" ", "").lower() in text.replace(" ", "").lower():
                print(f"[~] Found (ignore space): '{frag}'")
            else:
                print(f"[ ] MISSING: '{frag}'")

if __name__ == "__main__":
    test_full_extraction()
