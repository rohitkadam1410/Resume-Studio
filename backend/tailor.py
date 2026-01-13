import openai
import os
import json
import shutil
from docx import Document
from typing import List, Dict
import logging
import mlflow
from prompts import ANALYZE_GAPS_PROMPT_TEMPLATE, CALCULATE_SCORES_PROMPT_TEMPLATE

logger = logging.getLogger(__name__)

# Configure MLflow
MLFLOW_DB_PATH = "sqlite:///mlflow.db"
mlflow.set_tracking_uri(MLFLOW_DB_PATH)
mlflow.set_experiment("Resume Tailor Analysis")

# Ensure API key is set
# Ensure API key is set
# openai.api_key = os.environ.get("OPENAI_API_KEY")

from pydantic import BaseModel

class Edit(BaseModel):
    target_text: str
    new_content: str
    action: str
    rationale: str

class SectionAnalysis(BaseModel):
    section_name: str
    section_type: str = "Other"
    gaps: List[str]
    suggestions: List[str]
    edits: List[Edit]

class AnalysisResult(BaseModel):
    sections: List[SectionAnalysis]
    initial_score: int
    projected_score: int
    company_name: str = "Unknown Company"
    job_title: str = "Unknown Role"
    score_reasoning: str = ""



import re

def normalize_text(text: str) -> str:
    """Normalize whitespace for robust matching."""
    return re.sub(r'\s+', ' ', text).strip()

def safe_replace_text(paragraph, target: str, replacement: str):
    """
    Attempts to replace 'target' with 'replacement' in the paragraph while preserving formatting.
    Uses normalized matching to ignore whitespace differences from PDF conversion.
    """
    norm_target = normalize_text(target)
    norm_para_text = normalize_text(paragraph.text)

    if norm_target not in norm_para_text:
        return False

    # Check for multi-line or bulleted content
    # If the LLM returned markdown bullets, we want to strip them because 
    # pasting "* Item" into a DOCX doesn't make it a bullet list, just text.
    # We will strip them to keep it clean.
    clean_replacement = replacement
    if '\n' in replacement or any(line.strip().startswith(('* ', '- ', '• ')) for line in replacement.split('\n')):
         clean_replacement = re.sub(r'(^|\n)[\*\-\•]\s+', r'\1', replacement)

    # Use the cleaned replacement for the actual operation
    final_replacement = clean_replacement

    # 1. Try single run replacement (exact match first)
    for run in paragraph.runs:
        if target in run.text:
            run.text = run.text.replace(target, final_replacement)
            return True

    # 2. Relaxed match: Check if normalized target is in normalized run text
    for run in paragraph.runs:
        if norm_target in normalize_text(run.text):
            pattern = re.escape(target).replace(r'\ ', r'\s+')
            run.text = re.sub(pattern, final_replacement, run.text)
            return True

    # 3. Fallback: Full paragraph reconstruction
    style_run = paragraph.runs[0] if paragraph.runs else None
    font_name = style_run.font.name if style_run else None
    font_size = style_run.font.size if style_run else None
    bold = style_run.bold if style_run else None
    italic = style_run.italic if style_run else None
    
    pattern = re.escape(target).replace(r'\ ', r'\s+')
    new_text = re.sub(pattern, final_replacement, paragraph.text)
    
    if new_text != paragraph.text:
        paragraph.text = new_text
        for run in paragraph.runs:
            if font_name: run.font.name = font_name
            if font_size: run.font.size = font_size
            if bold is not None: run.bold = bold
            if italic is not None: run.italic = italic
        return True

    return False

def apply_edits_to_docx(docx_path: str, edits: List[Dict[str, str]], output_path: str):
    # Copy original to output
    shutil.copy2(docx_path, output_path)
    doc = Document(output_path)
    
    for edit in edits:
        if isinstance(edit, dict):
            action = edit.get("action")
            target = edit.get("target_text")
            content = edit.get("new_content")
        else:
            action = edit.action
            target = edit.target_text
            content = edit.new_content
        
        if not target or not content:
            continue
            
        def process_para(para):
            # We check if the normalized target is in the normalized paragraph text
            if normalize_text(target) in normalize_text(para.text):
                if action == "replace":
                    return safe_replace_text(para, target, content)
                elif action == "append":
                     para.add_run(" " + content)
                     return True
            return False

        # 1. Check body paragraphs
        found = False
        for para in doc.paragraphs:
            if process_para(para):
                found = True
                # Continue searching because the same target might appear multiple times
                # (e.g. name, contact info)
        
        # 2. Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if process_para(para):
                            found = True

    doc.save(output_path)

def extract_text_from_docx(docx_path: str) -> str:
    from docx.oxml.ns import qn
    doc = Document(docx_path)
    full_text = []

    # Track processed elements using the underlying XML element (stable identity)
    processed_elements = set()

    def add_to_full_text(elements, source_label=""):
        """Process a list of paragraph/table elements and add unique ones to full_text."""
        for element in elements:
            # For paragraphs, we use the _element (stable XML object)
            # For tables, they also have _element
            element_id = element._element if hasattr(element, '_element') else element
            
            if element_id in processed_elements:
                continue
            processed_elements.add(element_id)

            if hasattr(element, 'text'):
                if element.text.strip():
                    full_text.append(element.text)
            elif hasattr(element, 'rows'): # It's a table
                for row in element.rows:
                    for cell in row.cells:
                        # Cell objects are proxies, use the underlying tc (table cell) element
                        tc_id = cell._element
                        if tc_id not in processed_elements:
                            processed_elements.add(tc_id)
                            # Combine paragraphs in cell with space to avoid word merging
                            cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                            if cell_text.strip():
                                full_text.append(cell_text)

    # 1. Headers
    for section in doc.sections:
        add_to_full_text(section.header.paragraphs)
        add_to_full_text(section.header.tables)

    # 2. Body
    add_to_full_text(doc.paragraphs)
    add_to_full_text(doc.tables)

    # 3. Footers
    for section in doc.sections:
        add_to_full_text(section.footer.paragraphs)
        add_to_full_text(section.footer.tables)

    # 4. Textboxes (w:txbxContent)
    for element in doc.element.body.iter():
       if element.tag.endswith('txbxContent'):
           for p in element.findall(qn('w:p')):
               text = ""
               for r in p.findall(qn('w:r')):
                   for t in r.findall(qn('w:t')):
                       if t.text: text += t.text
               if text.strip():
                   # Basic text cleaning and check for textboxes
                   full_text.append(text)

    return '\n'.join(full_text)


# Removed extract_text_from_pdf dependency to ensure Sync

def analyze_gaps(docx_path: str, job_description: str, pdf_path: str = None) -> Dict:
    # Reverting to DOCX extraction to ensure identifying target_text works for replacement.
    # We improved extract_text_from_docx to include textboxes/tables.
    resume_text = extract_text_from_docx(docx_path)
    
    if not resume_text.strip():
        logger.warning("Extracted text is empty.")
    
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

    prompt = ANALYZE_GAPS_PROMPT_TEMPLATE.format(
        job_description=job_description,
        resume_text=resume_text
    )
    
    with mlflow.start_run(run_name="analyze_gaps"):
        mlflow.log_param("model", "gpt-4o")
        # Store prompt in DB via Tags (limit 5000 chars)
        mlflow.set_tag("prompt_template", ANALYZE_GAPS_PROMPT_TEMPLATE[:5000])
        mlflow.log_param("jd_length", len(job_description))
        mlflow.log_param("resume_length", len(resume_text))
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            response_format={ "type": "json_object" }
        )
        
        try:
            result = json.loads(response.choices[0].message.content)
            # Ensure extracted metadata exists
            result.setdefault("company_name", "Unknown Company")
            result.setdefault("job_title", "Unknown Role")
            
            # Log some basic metrics if available
            if "sections" in result:
                mlflow.log_metric("num_sections", len(result["sections"]))
            
        except json.JSONDecodeError:
            logger.error("Failed to decode JSON from LLM")
            result = {"sections": [], "company_name": "Unknown", "job_title": "Unknown"}
            mlflow.log_param("error", "JSONDecodeError")

        # 2. Separate Robust Scoring Step
        try:
            # Summarize proposed changes for the scorer
            changes_summary = []
            if "sections" in result:
                for section in result["sections"]:
                    name = section.get("section_name", "Unknown")
                    gaps = section.get("gaps", [])
                    suggestions = section.get("suggestions", [])
                    edits = section.get("edits", [])
                    changes_summary.append(f"Section {name}: Found {len(gaps)} gaps. {len(suggestions)} suggestions. Suggested {len(edits)} edits.")
                    if suggestions:
                        changes_summary.append(f" - Advice: {'; '.join(suggestions[:3])}")
            
            changes_text = "\n".join(changes_summary)
            scores = calculate_scores(resume_text, job_description, changes_text)
            result["initial_score"] = scores.get("initial_score", 0)
            result["projected_score"] = scores.get("projected_score", 0)
            result["score_reasoning"] = scores.get("reasoning", "")
            
            mlflow.log_metric("initial_score", result["initial_score"])
            mlflow.log_metric("projected_score", result["projected_score"])
            
        except Exception as e:
            logger.error(f"Scoring failed: {e}")
            result["initial_score"] = 0
            result["projected_score"] = 0
            mlflow.log_param("scoring_error", str(e))

        return result

def calculate_scores(resume_text: str, job_description: str, changes_summary: str) -> Dict:
    client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    
    prompt = CALCULATE_SCORES_PROMPT_TEMPLATE.format(
        job_description=job_description[:2000], # Truncated in original too, keeping consistent but cleaner
        resume_text=resume_text[:3000],
        changes_summary=changes_summary
    )
    
    # We create a nested run or a separate run? Usually scoring is part of the parent task.
    # For simplicity, we'll just log this as part of the parent run if it's active.
    # But calculate_scores might be called independently? The current flow calls it FROM analyze_gaps.
    # To be safe, checking active run.
    
    active_run = mlflow.active_run()
    if not active_run:
        # If called independently, start a run.
        # But here it's helper. Let's just log params to active run if exists.
        pass
    
    # Note: If we want to log the SCORING prompt specifically, we can log it as a param too.
    if active_run:
        mlflow.set_tag("scoring_prompt_template", CALCULATE_SCORES_PROMPT_TEMPLATE[:5000])

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            response_format={ "type": "json_object" }
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        logger.error(f"Error in calculate_scores: {e}")
        return {"initial_score": 0, "projected_score": 0}

def generate_tailored_resume(docx_path: str, sections: List[Dict]) -> str:
    # Flatten edits from all sections
    all_edits = []
    
    # Handle if sections is actually the AnalysisResult dict or list
    iterable_sections = sections if isinstance(sections, list) else sections.get("sections", [])
    
    for section in iterable_sections:
        # Handle dict vs object
        if isinstance(section, dict):
            if "edits" in section:
                all_edits.extend(section["edits"])
        else:
             all_edits.extend(section.edits)
            
    output_path = docx_path.replace(".docx", "_tailored.docx")
    apply_edits_to_docx(docx_path, all_edits, output_path)
    return output_path

