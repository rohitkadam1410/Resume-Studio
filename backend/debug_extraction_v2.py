import sys
import os
from docx import Document
from docx.oxml.ns import qn

def extract_text_from_docx(docx_path: str) -> str:
    print(f"Extracting from: {docx_path}")
    doc = Document(docx_path)
    full_text = []
    processed_elements = set()

    def add_to_full_text(elements, source_label=""):
        count = 0
        for element in elements:
            element_id = element._element if hasattr(element, '_element') else element
            if element_id in processed_elements:
                continue
            processed_elements.add(element_id)

            if hasattr(element, 'text'):
                if element.text.strip():
                    full_text.append(f"[{source_label}] {element.text}")
                    count += 1
            elif hasattr(element, 'rows'): 
                for row in element.rows:
                    for cell in row.cells:
                        tc_id = cell._element
                        if tc_id not in processed_elements:
                            processed_elements.add(tc_id)
                            # Simple join for debug
                            cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                            if cell_text.strip():
                                full_text.append(f"[{source_label} Table] {cell_text}")
                                count += 1
        return count

    # 1. Headers
    print("Checking Headers...")
    for section in doc.sections:
        add_to_full_text(section.header.paragraphs, "Header")
        add_to_full_text(section.header.tables, "Header")

    # 2. Body
    print("Checking Body...")
    add_to_full_text(doc.paragraphs, "Body")
    add_to_full_text(doc.tables, "Body")

    # 3. Textboxes
    print("Checking Textboxes...")
    for element in doc.element.body.iter():
       if element.tag.endswith('txbxContent'):
           for p in element.findall(qn('w:p')):
               text = ""
               for r in p.findall(qn('w:r')):
                   for t in r.findall(qn('w:t')):
                       if t.text: text += t.text
               if text.strip():
                   full_text.append(f"[Textbox] {text}")

    return '\n'.join(full_text)

if __name__ == "__main__":
    # Find a resume file
    target_file = None
    for file in os.listdir("d:/projects/LLM-Projects/Resume-Tailor/backend"):
        if file.startswith("temp_") and file.endswith(".docx") and "Kadam" in file:
            target_file = os.path.join("d:/projects/LLM-Projects/Resume-Tailor/backend", file)
            break
    
    if target_file:
        text = extract_text_from_docx(target_file)
        
        with open("backend/debug_extraction_output.txt", "w", encoding="utf-8") as f:
            f.write(text)
            
        print("Written to backend/debug_extraction_output.txt")
        
        # Check for summary keywords
        lower_text = text.lower()
        print(f"Contains 'professional summary': {'professional summary' in lower_text}")
        print(f"Contains 'summary': {'summary' in lower_text}")
        print(f"Contains 'profile': {'profile' in lower_text}")
    else:
        print("No temp resume found.")
