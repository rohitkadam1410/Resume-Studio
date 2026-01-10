from docx import Document

def create_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Developer at Tech Corp (2020-Present)')
    doc.add_paragraph('Implemented cool stuff using Python and React.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.Sc Computer Science')
    
from docx2pdf import convert
import pythoncom

def create_resume():
    doc = Document()
    doc.add_heading('John Doe', 0)
    doc.add_paragraph('Software Engineer')
    
    doc.add_heading('Experience', level=1)
    doc.add_paragraph('Senior Developer at Tech Corp (2020-Present)')
    doc.add_paragraph('Implemented cool stuff using Python and React.')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('B.Sc Computer Science')
    
    docx_path = 'd:/projects/LLM-Projects/Resume-Tailor/test_resume.docx'
    doc.save(docx_path)
    print("Created test_resume.docx")
    
    pythoncom.CoInitialize()
    convert(docx_path)
    print("Created test_resume.pdf")

if __name__ == "__main__":
    create_resume()
