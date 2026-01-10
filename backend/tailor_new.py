from docx.oxml.ns import qn
from docx import Document

def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    full_text = []

    # Track processed elements to avoid duplicates (especially in merged cells)
    processed_elements = set()

    def get_cell_id(cell):
        # Use the memory address of the internal XML element as a proxy for identity
        return id(cell._element)

    # 1. Headers (often contain contact info)
    for section in doc.sections:
        # Header Paragraphs
        for para in section.header.paragraphs:
            if para.text.strip() and id(para._element) not in processed_elements:
                full_text.append(para.text)
                processed_elements.add(id(para._element))
        
        # Header Tables
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if get_cell_id(cell) in processed_elements:
                        continue
                    processed_elements.add(get_cell_id(cell))
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)

    # 2. Body Paragraphs
    for para in doc.paragraphs:
        if para.text.strip() and id(para._element) not in processed_elements:
            full_text.append(para.text)
            processed_elements.add(id(para._element))
            
    # 3. Body Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Content of merged cells is repeated in each cell of the range
                # We track unique cells to print content only once
                if get_cell_id(cell) in processed_elements:
                    continue
                processed_elements.add(get_cell_id(cell))
                
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text)

    # 4. Textboxes (w:txbxContent)
    for element in doc.element.body.iter():
       if element.tag.endswith('txbxContent'):
           for p in element.findall(qn('w:p')):
               text = ""
               for r in p.findall(qn('w:r')):
                   for t in r.findall(qn('w:t')):
                       if t.text: text += t.text
               if text.strip():
                   full_text.append(text)

    return '\n'.join(full_text)
