from docx import Document
from docx.oxml.ns import qn

def analyze_table_structure(docx_path):
    doc = Document(docx_path)
    
    output = []
    output.append(f"=== Analyzing: {docx_path} ===\n")
    
    for table_idx, table in enumerate(doc.tables):
        output.append(f"\n--- TABLE {table_idx + 1} ---")
        for row_idx, row in enumerate(table.rows):
            output.append(f"  Row {row_idx + 1}:")
            for cell_idx, cell in enumerate(row.cells):
                cell_id = id(cell._element)
                para_count = len(cell.paragraphs)
                text_preview = " ".join([p.text[:50] for p in cell.paragraphs if p.text.strip()])[:100]
                output.append(f"    Cell {cell_idx + 1} (ID:{cell_id}, Paras:{para_count}): {text_preview}...")
    
    return "\n".join(output)

# Find temp file
import os
files = [f for f in os.listdir(".") if f.endswith(".docx") and f.startswith("temp_") and not f.endswith("_tailored.docx")]
if files:
    result = analyze_table_structure(files[0])
    with open("table_structure_analysis.txt", "w", encoding="utf-8") as f:
        f.write(result)
    print("Analysis written to table_structure_analysis.txt")
else:
    print("No temp files found")
